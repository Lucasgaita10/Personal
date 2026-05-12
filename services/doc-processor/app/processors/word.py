"""DOCX processor."""
from __future__ import annotations

import io
from docx import Document


def extract_docx(buf: bytes) -> str:
    d = Document(io.BytesIO(buf))
    parts: list[str] = []
    for p in d.paragraphs:
        if p.text:
            parts.append(p.text)
    for table in d.tables:
        for row in table.rows:
            parts.append("\t".join(c.text for c in row.cells))
    return "\n".join(parts)
